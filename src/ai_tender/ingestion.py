import hashlib
from collections.abc import Iterable
from pathlib import Path

import fitz
from docx import Document
from openpyxl import load_workbook

from .models import Block

SUPPORTED = {".docx", ".pdf", ".xlsx"}


def _block(path: Path, location: str, kind: str, text: str) -> Block | None:
    text = " ".join(text.split())
    if len(text) < 20:
        return None
    key = f"{path.resolve()}:{location}:{text[:100]}".encode("utf-8")
    return Block(
        id=hashlib.sha256(key).hexdigest()[:16],
        file=str(path),
        location=location,
        kind=kind,
        text=text,
    )


def _chunks(text: str, size: int = 5000) -> Iterable[str]:
    paragraphs = [part.strip() for part in text.splitlines() if part.strip()]
    current: list[str] = []
    length = 0
    for paragraph in paragraphs:
        if current and length + len(paragraph) > size:
            yield "\n".join(current)
            current, length = [], 0
        current.append(paragraph)
        length += len(paragraph)
    if current:
        yield "\n".join(current)


def extract_docx(path: Path) -> list[Block]:
    document = Document(path)
    blocks: list[Block] = []
    paragraph_group: list[str] = []
    group_number = 1
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        is_heading = paragraph.style and paragraph.style.name.lower().startswith("heading")
        if is_heading and paragraph_group:
            item = _block(path, f"абзацы {group_number}", "text", "\n".join(paragraph_group))
            if item:
                blocks.append(item)
            paragraph_group, group_number = [], group_number + 1
        paragraph_group.append(text)
    if paragraph_group:
        for index, text in enumerate(_chunks("\n".join(paragraph_group)), start=group_number):
            item = _block(path, f"текстовый блок {index}", "text", text)
            if item:
                blocks.append(item)

    for index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        item = _block(path, f"таблица {index}", "table", "\n".join(rows))
        if item:
            blocks.append(item)
    return blocks


def extract_pdf(path: Path) -> tuple[list[Block], list[str]]:
    blocks: list[Block] = []
    warnings: list[str] = []
    with fitz.open(path) as document:
        for page_number, page in enumerate(document, start=1):
            text = page.get_text("text").strip()
            if not text:
                warnings.append(f"Нет текстового слоя: {path}, стр. {page_number}; требуется OCR")
                continue
            for part_number, text_part in enumerate(_chunks(text), start=1):
                item = _block(
                    path,
                    f"стр. {page_number}, блок {part_number}",
                    "text",
                    text_part,
                )
                if item:
                    blocks.append(item)
    return blocks, warnings


def extract_xlsx(path: Path) -> list[Block]:
    workbook = load_workbook(path, read_only=True, data_only=True)
    blocks: list[Block] = []
    for sheet in workbook.worksheets:
        rows: list[str] = []
        part = 1
        for row in sheet.iter_rows(values_only=True):
            cells = [str(value).strip() if value is not None else "" for value in row]
            if any(cells):
                rows.append(" | ".join(cells))
            if len(rows) >= 100:
                item = _block(path, f"лист {sheet.title}, блок {part}", "table", "\n".join(rows))
                if item:
                    blocks.append(item)
                rows, part = [], part + 1
        if rows:
            item = _block(path, f"лист {sheet.title}, блок {part}", "table", "\n".join(rows))
            if item:
                blocks.append(item)
    workbook.close()
    return blocks


def extract_folder(folder: Path, technical_only: bool = False) -> tuple[list[Block], list[str]]:
    if not folder.is_dir():
        raise ValueError(f"Папка не найдена: {folder}")
    files = [path for path in folder.rglob("*") if path.is_file()]
    if technical_only:
        preferred = [
            path
            for path in files
            if any(marker in str(path).lower() for marker in ("тз", "техническ", "извещение"))
        ]
        if preferred:
            files = preferred

    blocks: list[Block] = []
    warnings: list[str] = []
    for path in files:
        suffix = path.suffix.lower()
        try:
            if suffix == ".docx":
                blocks.extend(extract_docx(path))
            elif suffix == ".pdf":
                extracted, pdf_warnings = extract_pdf(path)
                blocks.extend(extracted)
                warnings.extend(pdf_warnings)
            elif suffix == ".xlsx":
                blocks.extend(extract_xlsx(path))
            elif suffix in {".doc", ".rar"}:
                warnings.append(f"Формат пока пропущен ({suffix}): {path}")
        except Exception as exc:
            warnings.append(f"Ошибка чтения {path}: {exc}")
    return blocks, warnings
