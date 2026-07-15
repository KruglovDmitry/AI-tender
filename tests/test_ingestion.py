from pathlib import Path

from docx import Document

from ai_tender.ingestion import extract_folder


def test_extract_docx_paragraphs_and_tables(tmp_path: Path) -> None:
    path = tmp_path / "Техническое задание.docx"
    document = Document()
    document.add_paragraph("Требуемый класс точности счётчика должен быть не хуже 1.0.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Степень защиты"
    table.cell(1, 1).text = "IP54"
    document.save(path)

    blocks, warnings = extract_folder(tmp_path, technical_only=True)

    assert not warnings
    assert any("класс точности" in block.text for block in blocks)
    assert any(block.kind == "table" and "IP54" in block.text for block in blocks)


def test_missing_folder_is_rejected(tmp_path: Path) -> None:
    missing = tmp_path / "missing"

    try:
        extract_folder(missing)
    except ValueError as exc:
        assert "Папка не найдена" in str(exc)
    else:
        raise AssertionError("Ожидалась ошибка для отсутствующей папки")
