import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document

from ai_tender.services.archive_service import (
    configure_rarfile,
    expand_archives,
    find_unrar_tool,
    unpack_zip,
)
from ai_tender.services.loader_service import is_ignored_file, load_documents
from ai_tender.services.ocr_service import ocr_status


def test_find_unrar_tool_prefers_env(tmp_path: Path, monkeypatch) -> None:
    fake = tmp_path / "unrar.exe"
    fake.write_bytes(b"stub")
    monkeypatch.setenv("UNRAR_TOOL", str(fake))
    assert find_unrar_tool() == fake


def test_configure_rarfile_sets_tool(tmp_path: Path, monkeypatch) -> None:
    fake = tmp_path / "UnRAR.exe"
    fake.write_bytes(b"stub")
    monkeypatch.setenv("UNRAR_TOOL", str(fake))
    assert configure_rarfile() == fake

    import rarfile

    assert rarfile.UNRAR_TOOL == str(fake)


def test_ocr_status_reports_missing_tesseract(monkeypatch) -> None:
    monkeypatch.delenv("TESSERACT_CMD", raising=False)
    monkeypatch.delenv("TESSERACT_PATH", raising=False)
    with patch("ai_tender.services.ocr_service.tesseract_path", return_value=None):
        ok, message = ocr_status()
    assert not ok
    assert "Tesseract" in message


def test_load_documents_reads_docx(tmp_path: Path) -> None:
    path = tmp_path / "Техническое задание.docx"
    document = Document()
    document.add_paragraph("Требуемый класс точности счётчика должен быть не хуже 1.0.")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Параметр"
    table.cell(0, 1).text = "Значение"
    table.cell(1, 0).text = "Степень защиты"
    table.cell(1, 1).text = "IP54"
    document.save(path)

    docs, warnings = load_documents(tmp_path, corpus="tender", technical_only=True)

    assert not warnings
    text = " ".join(doc.text for doc in docs)
    assert "класс точности" in text
    assert "IP54" in text


def test_ignored_files_are_not_warned(tmp_path: Path) -> None:
    (tmp_path / "Thumbs.db").write_bytes(b"junk")
    (tmp_path / "нормальный.txt").write_text("достаточно длинный текст документа", encoding="utf-8")

    docs, warnings = load_documents(tmp_path, corpus="tender")

    assert any("достаточно длинный" in doc.text for doc in docs)
    assert not any("Thumbs.db" in warning for warning in warnings)


def test_load_documents_reads_doc_with_mock(tmp_path: Path, monkeypatch) -> None:
    path = tmp_path / "протокол.doc"
    path.write_bytes(b"fake-doc-bytes")

    def fake_extract_doc_text(file_path: Path, **kwargs: object) -> tuple[str, str | None]:
        return "Требование из legacy doc: класс точности 1.0.", None

    monkeypatch.setattr("ai_tender.services.loader_service.extract_doc_text", fake_extract_doc_text)

    docs, warnings = load_documents(tmp_path, corpus="tender")

    assert not warnings
    assert any("класс точности" in doc.text for doc in docs)
    assert docs[0].metadata["file_path"] == "протокол.doc"


def test_is_ignored_file() -> None:
    assert is_ignored_file("архив/Thumbs.db")
    assert is_ignored_file("~$документ.docx")
    assert not is_ignored_file("ТЗ/протокол.doc")


def test_missing_folder_is_rejected(tmp_path: Path) -> None:
    missing = tmp_path / "missing"

    try:
        load_documents(missing, corpus="tender")
    except ValueError as exc:
        assert "Папка не найдена" in str(exc)
    else:
        raise AssertionError("Ожидалась ошибка для отсутствующей папки")


def test_load_documents_reads_docx_inside_zip(tmp_path: Path) -> None:
    source = tmp_path / "Техническое задание.docx"
    document = Document()
    document.add_paragraph("Требование из архива: степень защиты не ниже IP54.")
    document.save(source)

    archive = tmp_path / "пакет.zip"
    with zipfile.ZipFile(archive, "w") as handle:
        handle.write(source, arcname="docs/Техническое задание.docx")
    source.unlink()

    docs, warnings = load_documents(tmp_path, corpus="tender", technical_only=True)

    assert not warnings
    assert any("IP54" in doc.text for doc in docs)
    assert any("пакет.zip/docs/Техническое задание.docx" in doc.metadata["file_path"] for doc in docs)


def test_zip_slip_is_rejected(tmp_path: Path) -> None:
    archive = tmp_path / "evil.zip"
    with zipfile.ZipFile(archive, "w") as handle:
        handle.writestr("../escape.txt", "secret")

    dest = tmp_path / "out"
    dest.mkdir()
    try:
        unpack_zip(archive, dest)
    except ValueError as exc:
        assert "Небезопасный путь" in str(exc)
    else:
        assert not (tmp_path / "escape.txt").exists()


def test_expand_archives_preserves_nested_paths(tmp_path: Path) -> None:
    warnings: list[str] = []
    temp_dirs = []
    nested = tmp_path / "nested.txt"
    nested.write_text("nested content", encoding="utf-8")
    archive = tmp_path / "pack.zip"
    with zipfile.ZipFile(archive, "w") as handle:
        handle.write(nested, arcname="docs/nested.txt")
    nested.unlink()

    items = expand_archives([(archive, "pack.zip")], temp_dirs, warnings)
    for temp in temp_dirs:
        temp.cleanup()

    assert any(label == "pack.zip/docs/nested.txt" for _, label in items)
